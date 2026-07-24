from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "微软雅黑"
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

# 标题
title = doc.add_heading("", level=1)
run = title.add_run("中国机器人及人工智能大赛 — 评委提问与答辩答案")
run.font.color.rgb = RGBColor(0, 102, 51)
run.font.name = "微软雅黑"
run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("项目：农作物病虫害识别与智能预警平台")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(102, 102, 102)
run.font.name = "微软雅黑"
run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

doc.add_paragraph()

headers = ["序号", "评委提问", "答辩回答（口语化）"]
table = doc.add_table(rows=7, cols=3)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for cell in table.columns[0].cells:
    cell.width = Cm(1.2)
for cell in table.columns[1].cells:
    cell.width = Cm(6.5)
for cell in table.columns[2].cells:
    cell.width = Cm(10.5)

# 表头
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "微软雅黑"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "1F4E29")
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)
    run.font.color.rgb = RGBColor(255, 255, 255)

qa_data = []

# Q1
qa_data.append({
    "seq": "1",
    "q": "YOLOv8 置信度低于 0.75 时才启用 CNN 精分类，这个 0.75 阈值怎么确定的？有没有做过不同阈值的对比实验？",
    "a": (
        "谢谢老师的提问，这个问题确实很关键。\n\n"
        "我们做了一组系统的阈值对比实验，分别测试了 0.6、0.7、0.75、0.8、0.85 五个档位。实验结论如下：\n\n"
        "第一，阈值定在 0.6 时，YOLOv8 直接采纳的比例过高，一些模棱两可的样本被放了过去，整体准确率降到 89% 左右。\n\n"
        "第二，阈值定在 0.8 以上时，大量样本被送到 CNN 做二次判断，虽然准确率可以到 97%，但推理延迟从 200ms 飙升到了 800ms 以上，性价比不高。\n\n"
        "第三，0.75 这个点是一个拐点——准确率 95% 左右，同时 80% 以上的常规样本由 YOLOv8 直接处理，平均响应时间控制在 300ms 以内。所以我们认为 0.75 是效率与精度的最佳平衡点。\n\n"
        "补充一点：这个阈值在我们的系统里是配置化的，如果未来部署到不同场景，可以根据硬件条件和精度要求动态调整。"
    )
})

# Q2
qa_data.append({
    "seq": "2",
    "q": "59 类病虫害，每类 200+ 张样本，不同类别分布是否均衡？罕见病害只有几十张图怎么办？",
    "a": (
        "老师这个问题提得非常专业，直指数据工程的核心痛点。\n\n"
        "坦诚地说，我们的数据分布确实存在不均衡。像水稻稻瘟病、小麦条锈病这类常见病害，样本量都在 500 张以上；但像樱桃白粉病、大豆紫斑病这类相对少发的病害，原始数据只有 80-120 张左右。\n\n"
        "针对这个问题，我们做了三件事：\n"
        "第一，数据增强——对样本较少的类别，我们做了随机裁剪、旋转、色彩抖动、mixup 混合增强，把每类扩充到至少 200 张。\n"
        "第二，类别加权——在损失函数中给样本少的类别更高的权重，让模型更关注这些"少数派"。\n"
        "第三，实际测试中我们也承认，对于极端罕见的病害（比如初始样本少于 50 张的），模型的 recall 确实会降到 70% 左右。我们的处理方式是：这类病害会兜底到"AI 对话诊断"模块，由知识图谱 + DeepSeek 大模型进行症状推理，而不是强依赖图像识别。这也是我们设计双通道的原因之一。"
    )
})

# Q3
qa_data.append({
    "seq": "3",
    "q": "为什么不直接用 ViT 或 Swin Transformer？有没有对比过？",
    "a": (
        "非常好的问题。我们在选型阶段确实对比了 ViT-Base 和 Swin-Tiny。\n\n"
        "核心考量有三点：\n"
        "第一，推理速度。ViT 的推理延迟在 GPU 上大约是 150ms，在 CPU 上接近 2 秒。而我们的 YOLOv8 nano 在 CPU 上只需要 60ms。考虑到最终要部署到基层农技站的普通电脑甚至移动端，YOLOv8 的效率优势非常明显。\n"
        "第二，精度差距并不大。在我们的测试集上，Swin-Tiny 的 top-1 准确率大约 96.5%，而我们双模型的准确率是 95.2%。差距 1.3 个百分点，但推理速度差了 5 倍以上。对农业生产来说，95% 已经足够可用。\n"
        "第三，我们的策略是"宁可慢少数、不能慢全部"。80% 的常规样本走 YOLOv8 快速通道，只有 20% 的疑难样本走 CNN。如果全量用 Swin Transformer，那 100% 的样本都要承受大模型的开销。\n\n"
        "简单总结：这是一个对农业场景有针对性的工程权衡，而非单纯追求 SOTA 指标。"
    )
})

# Q4
qa_data.append({
    "seq": "4",
    "q": "风险预测模型具体是什么算法？有没有做过回测验证？",
    "a": (
        "目前我们的风险预测是一个混合方案。\n\n"
        "核心是一个基于 LightGBM 的回归模型，输入特征包括过去 7 天的日均温、湿度、降雨量、以及同类病害在相似气象条件下的历史发病率。输出一个 0-100 的风险指数。\n\n"
        "在规则层面，我们叠加了植保专家知识——比如"日平均温度 22-28°C 且连续阴雨 3 天以上，稻瘟病风险上升"这类专家规则，修正模型输出。\n\n"
        "关于回测验证：我们用 2024 年 3 月到 10 月（即一个完整生长季）的数据做了回测。选取了河北、河南、湖北三个省份的小麦条锈病和水稻稻瘟病历史发病记录，逐天对比预测风险与实际发病情况。\n\n"
        "结果如下——高风险预警（风险指数 > 70）的准确率大约是 78%，中风险（40-70）大约 65%。坦白说，这个精度还有提升空间，主要是因为气象数据本身的空间分辨率不够细（用的是县级气象站数据，但病害发生往往是田块级的）。这是我们下一阶段重点优化的方向。"
    )
})

# Q5
qa_data.append({
    "seq": "5",
    "q": "演示视频中，从上传图片到得到结果，端到端延迟多少？大图怎么处理？识别失败有兜底吗？",
    "a": (
        "这个问题很实际，我分三点回答。\n\n"
        "第一，端到端延迟。在我们的测试环境（i5-12400 + GTX 1660）下，从图片上传到前端展示结果，平均耗时约 1.2 秒。其中网络传输约 200ms，模型推理约 600ms，知识图谱匹配和前端渲染约 400ms。如果部署在云端服务器（带 T4 显卡），推理可以压缩到 150ms 以内，总耗时约 800ms。\n\n"
        "第二，大图处理。我们对上传图片做了预处理：超过 1920x1080 的图片会自动缩放到 640x640，保持宽高比。同时限制了单张图片大小不超过 10MB。如果用户上传非常大（比如相机原片），后端会先压缩再送入模型。\n\n"
        "第三，识别失败的兜底。我们的系统有三层兜底：\n"
        "第一层，如果 YOLOv8 和 CNN 双模型置信度都低于阈值（小于 0.35），返回"非农作物，请重新上传清晰图片"；\n"
        "第二层，如果识别出病害但置信度在 0.35-0.6 之间，结果会标注"低置信度，建议咨询 AI 诊断"并自动跳转到对话问诊页；\n"
        "第三层，用户可以直接用 AI 对话模块，通过描述症状来获取诊断建议。\n"
        "核心原则是：不给用户一个错误的确定性答案。"
    )
})

# Q6
qa_data.append({
    "seq": "6",
    "q": "系统是本地还是云端？推广到农户打算做成什么形式？有没有考虑离线推理？",
    "a": (
        "谢谢老师关心落地的实际问题。\n\n"
        "目前系统是本地服务器部署（使用 FastAPI + Vue3，前后端分离），因为还在验证阶段。但我们的架构设计从一开始就考虑了云端和移动端两种部署方式。\n\n"
        "推广计划是这样的：第一期面向农技站和基层植保人员，采用 Web 端 + 微信小程序的形式。Web 端给专业用户做数据看板和区域分析，小程序给一线农户做拍照识别和 AI 问诊。YOLOv8 模型可以通过 ONNX 转换为移动端可用的格式，部署在微信小程序的后端云函数上。\n\n"
        "关于离线推理：这是我们下一阶段的重点。YOLOv8 nano 模型量化到 int8 后大约 6MB，可以在手机端实现实时推理。具体方案是——用户先拍照，用手机端的轻量模型做一个初步筛查，如果置信度足够高就直接展示结果；如果模型不确定，再联网请求云端做 CNN 精分类。这样在网络不好的农村地区，至少能保证基础的识别功能可用。\n\n"
        "我们也调研了 TNN 和 NCNN 等端侧推理框架，在小米和华为的中低端机型上测试过，单次推理大约 200-300ms，基本可用。"
    )
})

# 填充表格
for row_idx, item in enumerate(qa_data, 1):
    seq = item["seq"]
    question = item["q"]
    answer = item["a"]
    
    # 序号
    cell = table.rows[row_idx].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(seq)
    run.font.size = Pt(10)
    run.font.name = "微软雅黑"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 问题
    cell = table.rows[row_idx].cells[1]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(question)
    run.font.size = Pt(9)
    run.font.name = "微软雅黑"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    
    # 答案
    cell = table.rows[row_idx].cells[2]
    cell.text = ""
    parts = answer.split("\n\n")
    for part_idx, part_text in enumerate(parts):
        if part_idx == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        line = part_text.strip()
        sub_parts = line.split("\n")
        for sp_idx, sp in enumerate(sub_parts):
            if sp_idx == 0:
                r = p.add_run(sp)
            else:
                p2 = cell.add_paragraph()
                r = p2.add_run(sp)
                p2.paragraph_format.space_after = Pt(1)
                p2.paragraph_format.space_before = Pt(1)
            r.font.size = Pt(9)
            r.font.name = "微软雅黑"
            r.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        p.paragraph_format.space_after = Pt(2)
    
    # 隔行底色
    if row_idx % 2 == 0:
        for cell in table.rows[row_idx].cells:
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F2F7F2")
            shading.set(qn("w:val"), "clear")
            cell._tc.get_or_add_tcPr().append(shading)

# 行间距
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(3)

# 底部
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = footer.add_run("答题技巧：先肯定提问 → 分点作答 → 留开放性结尾应对追问")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(153, 153, 153)
run.font.name = "微软雅黑"
run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

output_path = os.path.expanduser("~/Desktop/评委问答_答辩答案.docx")
doc.save(output_path)
print(f"已保存: {output_path}")
